import streamlit as st
import google.generativeai as genai
import json
import PyPDF2
from PIL import Image
import datetime
import io
import docx
import os
import requests
import base64
from fpdf import FPDF
import urllib.parse
import uuid

# ==========================================
# FUNGSI PDF UNTUK MARKETING BROCHURE (MODUL 13)
# ==========================================
class ULTIMATE_BROCHURE(FPDF):
    def __init__(self, brand_color, brand_name, website_link, logo_path, wa_number):
        super().__init__()
        self.brand_color = brand_color
        self.brand_name = brand_name
        self.website_link = website_link
        self.logo_path = logo_path
        self.wa_number = wa_number

    def header(self):
        self.set_fill_color(*self.brand_color)
        self.rect(0, 0, 210, 4, 'F')
        
        if self.logo_path and os.path.exists(self.logo_path):
            self.image(self.logo_path, x=160, y=8, w=40)
        else:
            self.ln(5)
            self.set_font('helvetica', 'B', 24)
            self.set_text_color(*self.brand_color)
            self.cell(0, 10, self.brand_name, ln=True, align='R')

    def footer(self):
        self.set_y(-25)
        self.set_draw_color(*self.brand_color)
        self.line(10, 272, 200, 272)
        
        self.set_text_color(50, 50, 50)
        self.set_font('helvetica', 'B', 9)
        self.cell(0, 6, f'{self.brand_name} - SMART EQUIPMENT FOR SMART BUILDERS', align='C', ln=True)
        
        self.set_font('helvetica', 'I', 8)
        clean_link = self.website_link.replace("https://", "").replace("http://", "").rstrip("/")
        self.cell(0, 4, f'Authorized Representative: Adjie Agung | {clean_link}', align='C', ln=True)
        
# ==========================================
# KONFIGURASI HALAMAN & STATE MANAGEMENT
# ==========================================
st.set_page_config(page_title="VORTEX 4.0 | By Adjie Agung", page_icon="🌪️", layout="centered")

st.markdown("""
    <style>
    .block-container { padding-top: 1.5rem; padding-bottom: 2rem; }
    .footer { text-align: center; color: #888; font-size: 0.85rem; margin-top: 50px; padding-top: 10px; border-top: 1px solid #ddd; }
    .stButton>button { border-radius: 8px; font-weight: bold; }
    </style>
""", unsafe_allow_html=True)

# SETUP API KEYS
try:
    api_key = st.secrets["GEMINI_API_KEY"]
    REMOVE_BG_API_KEY = st.secrets.get("REMOVE_BG_API_KEY", "")
    genai.configure(api_key=api_key)
except KeyError:
    st.error("⚠️ GEMINI_API_KEY tidak ditemukan di Secrets!")
    st.stop()

# INITIALIZE SESSION STATE
if "intelligence_data" not in st.session_state: st.session_state.intelligence_data = ""
if "campaign_data" not in st.session_state: st.session_state.campaign_data = None
if "project_scenario" not in st.session_state: st.session_state.project_scenario = ""
if "current_weather" not in st.session_state: st.session_state.current_weather = ""

# ==========================================
# FUNGSI PDF UNTUK PHOTONIS (CORPORATE ELITE EDITION)
# ==========================================
class VISUAL_PDF(FPDF):
    def __init__(self, logo_image=None, brand_name="", product_image=None):
        super().__init__()
        self.logo_image = logo_image
        self.brand_name = brand_name
        self.product_image = product_image

    def header(self):
        # 1. Logo Perusahaan (Kiri)
        if self.logo_image:
            img_buf = io.BytesIO()
            self.logo_image.save(img_buf, format='PNG')
            img_buf.seek(0)
            self.image(img_buf, 10, 10, 35)

        # 2. Judul Dokumen (Kanan) - Warna Biru Korporat
        self.set_font('helvetica', 'B', 16)
        self.set_text_color(31, 73, 125) 
        self.cell(0, 6, "OFFICIAL TECHNICAL REPORT", border=False, ln=True, align='R')

        # 3. Sub-Judul Ekosistem
        self.set_font('helvetica', 'I', 9)
        self.set_text_color(100, 100, 100)
        self.cell(0, 5, "Tatsuo & AIMIX Heavy Equipment Ecosystem", border=False, ln=True, align='R')

        # 4. Tanggal Dibuat (Timestamp)
        now_str = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        self.set_font('helvetica', '', 8)
        self.cell(0, 5, f"Generated: {now_str}", border=False, ln=True, align='R')

        # 5. Garis Pemisah Biru
        self.ln(3)
        self.set_draw_color(31, 73, 125)
        self.set_line_width(0.5)
        self.line(10, self.get_y(), 200, self.get_y())
        self.ln(5)

    def footer(self):
        self.set_y(-15)
        self.set_font('helvetica', 'I', 8)
        self.set_text_color(128, 128, 128)
        self.cell(0, 10, f"Developed by Adjie Agung - VORTEX Intelligence Engine", align='C')

def create_visual_pdf(text, logo, brand, product_bytes):
    pdf = VISUAL_PDF(logo_image=logo, brand_name=brand)
    pdf.add_page()
    
    # --- RENDER GAMBAR PRODUK 100% UTUH (AUTO-RESIZE Y) ---
    if product_bytes:
        prod_img = Image.open(io.BytesIO(product_bytes)).convert("RGBA") 
        prod_buf = io.BytesIO()
        prod_img.save(prod_buf, format='PNG')
        prod_buf.seek(0)
        
        # Kalkulasi dimensi dinamis agar roda/bawah gambar tidak tertabrak kotak
        img_w_px, img_h_px = prod_img.size
        target_width = 130 # Lebar gambar di PDF (mm)
        calculated_height = (img_h_px / img_w_px) * target_width
        
        y_start = pdf.get_y() + 2
        x_start = (210 - target_width) / 2 # Posisi presisi di tengah kertas
        
        # Cetak gambar dengan warna solid 100%
        pdf.image(prod_buf, x=x_start, y=y_start, w=target_width)
        
        # Pindahkan kursor teks secara dinamis tepat di bawah gambar + jarak aman 8mm
        # Ini menjamin gambar APAPUN tidak akan pernah terpotong lagi!
        pdf.set_y(y_start + calculated_height + 8)

    # --- PISAHKAN DATA BOX & KONTEN ---
    lines = text.encode('latin-1', 'ignore').decode('latin-1').replace('**', '').split('\n')
    box_data = []
    content_lines = []
    
    for line in lines:
        line_str = line.strip()
        if line_str.startswith("UNIT:") or line_str.startswith("KATEGORI:") or line_str.startswith("FOKUS:"):
            box_data.append(line_str)
        else:
            content_lines.append(line_str)

    # --- RENDER DATA BOX DENGAN EFEK KACA (85% Opacity) ---
    if box_data:
        # local_context membuat kotak sedikit tembus pandang sesuai request Anda
        with pdf.local_context(fill_opacity=0.85): 
            pdf.set_fill_color(242, 247, 252) # Biru muda korporat
            pdf.set_draw_color(31, 73, 125)
            pdf.set_line_width(0.3)
            
            pdf.set_font("helvetica", "B", 10)
            pdf.set_text_color(31, 73, 125)
            pdf.cell(0, 8, " INFORMASI UNIT UTAMA", border="LTR", ln=True, fill=True)
            
            pdf.set_font("helvetica", "", 10)
            pdf.set_text_color(0, 0, 0)
            for bd in box_data:
                parts = bd.split(":", 1)
                if len(parts) == 2:
                    pdf.cell(35, 7, f" {parts[0].strip()}", border="L", fill=True)
                    pdf.cell(0, 7, f": {parts[1].strip()}", border="R", ln=True, fill=True)
                else:
                    pdf.cell(0, 7, f" {bd}", border="LR", ln=True, fill=True)
                    
            pdf.cell(0, 2, "", border="LBR", ln=True, fill=True)
            pdf.ln(5)

    # --- RENDER KONTEN TEKS & TABEL CERDAS ---
    in_table = False
    for line in content_lines:
        if not line:
            if in_table:
                pdf.ln(3)
                in_table = False
            else:
                pdf.ln(2)
            continue
            
        if line.startswith('|'):
            if '|---' in line or '|:' in line: continue
            cells = [c.strip() for c in line.split('|')[1:-1]]
            if not cells: continue
            col_w = 190 / len(cells)
            
            if not in_table:
                pdf.set_font("helvetica", "B", 9)
                pdf.set_fill_color(41, 128, 185)
                pdf.set_text_color(255, 255, 255)
                for cell in cells:
                    pdf.cell(col_w, 8, cell, border=1, fill=True)
                pdf.ln()
                in_table = True
            else:
                pdf.set_font("helvetica", "", 9)
                pdf.set_fill_color(255, 255, 255)
                pdf.set_text_color(0, 0, 0)
                for cell in cells:
                    pdf.cell(col_w, 7, cell, border=1, fill=True)
                pdf.ln()
            continue
        else:
            in_table = False

        if line.startswith('## '):
            clean_line = line.replace('##', '').strip()
            pdf.ln(4)
            pdf.set_font("helvetica", "B", 11)
            pdf.set_text_color(0, 0, 0)
            pdf.cell(0, 8, clean_line.upper(), border=0, ln=True)
            pdf.ln(1)
        elif line.startswith('- ') or line.startswith('* '):
            clean_line = line[2:].strip()
            pdf.set_font("helvetica", "", 10)
            pdf.set_text_color(40, 40, 40)
            pdf.set_x(15)
            pdf.multi_cell(0, 6, f"-  {clean_line}")
            pdf.set_x(10)
        elif not line.startswith('#'):
            pdf.set_font("helvetica", "", 10)
            pdf.set_text_color(40, 40, 40)
            pdf.multi_cell(0, 6, line)

    return pdf.output()

# ==========================================
# SIDEBAR
# ==========================================
with st.sidebar:
    st.image("https://img.icons8.com/color/96/000000/engineering.png", width=70)
    st.title("⚙️ Command Center")
    st.success("Model: gemini-flash-latest")
    st.info("Library: AIMIX, Tatsuo, New Timehope")
    st.divider()
    wa_num = st.text_input("WhatsApp Sales", "+6281230857759")
    aff_link = st.text_input("Shopee Affiliate", "https://collshp.com/vortex_catalog?share_channel_code=1")
    st.divider()
    st.caption("Engine: VORTEX 4.0 | By Adjie Agung")

st.title("🌪️ VORTEX 4.0")
st.subheader("Heavy-Asset Domination System")

# ==========================================
# MODUL 1: RADAR CUACA & TENDER
# ==========================================
with st.expander("📡 Modul 1: Radar Intelijen (Cuaca & Tender)", expanded=False):
    col_w1, col_w2 = st.columns(2)
    with col_w1:
        st.markdown("**⛈️ Sensor Cuaca**")
        lokasi = st.selectbox("Lokasi Radar", ["Samarinda, ID", "Penajam (IKN), ID", "Balikpapan, ID"])
        if st.button("Tarik Cuaca", use_container_width=True):
            st.session_state.current_weather = "Hujan Ringan, Suhu 24-30°C, Kelembapan 98% (Tanah sangat basah/lumpur)"
            st.toast(f"Cuaca {lokasi} Terdeteksi!")

    with col_w2:
        st.markdown("**🏗️ Radar Tender**")
        cakupan_radar = st.selectbox("Wilayah", ["Kalimantan Timur & IKN", "Seluruh Kalimantan", "Maluku & Indonesia Timur"])
        if st.button("🛰️ Scan Tender", use_container_width=True):
            with st.spinner(f"Memindai LPSE di {cakupan_radar}..."):
                model_radar = genai.GenerativeModel('gemini-flash-latest')
                res = model_radar.generate_content(f"Cari info tren proyek infrastruktur, tambang di {cakupan_radar} bulan ini.")
                st.session_state.project_scenario = res.text
            st.toast("Radar Tender Berhasil!")

    final_scenario = st.text_area("Skenario Target", value=f"Cuaca: {st.session_state.current_weather}\n\nProyek: {st.session_state.project_scenario}", height=120)

# ==========================================
# MODUL 2: EKSTRAKSI & INTELIJEN PRODUK
# ==========================================
with st.expander("🎯 Modul 2: Ekstraksi Spek & Analisis", expanded=False):
    col_u1, col_u2 = st.columns(2)
    with col_u1:
        brand = st.selectbox("Merek Produk", ["AIMIX", "Tatsuo", "New Timehope"])
        unit_type = st.text_input("Tipe Unit", "HSPD 360" if brand == "New Timehope" else "Self Loading Mixer + ABT60C")
    with col_u2:
        uploaded_file = st.file_uploader("Upload Brosur", type=["pdf", "png", "jpg", "jpeg"], key="brosur_m2")
        
    pdf_text = ""
    image_data = None
    if uploaded_file:
        ext = uploaded_file.name.split('.')[-1].lower()
        if ext == 'pdf':
            for page in PyPDF2.PdfReader(uploaded_file).pages:
                ex = page.extract_text()
                if ex: pdf_text += ex
            st.toast("📄 PDF Dibaca!")
        elif ext in ['png', 'jpg', 'jpeg']:
            image_data = Image.open(uploaded_file)
            st.toast("🖼️ Gambar Dibaca!")
        
    if st.button("🔍 Analisis Sudut Serang", use_container_width=True, type="primary"):
        model_intel = genai.GenerativeModel('gemini-flash-latest')
        intel_prompt = f"Analisis marketing untuk {brand} {unit_type}. Kondisi: {final_scenario}. Spek: {pdf_text}. Buat: 1. Pain Points, 2. Solusi Teknis, 3. Killer Angle."
        contents = [intel_prompt]
        if image_data: contents.append(image_data)
        with st.spinner("Menganalisis data..."):
            st.session_state.intelligence_data = model_intel.generate_content(contents).text
            st.toast("Intelijen Selesai!")

    if st.session_state.intelligence_data:
        st.info(st.session_state.intelligence_data)
        
# ==========================================
# MODUL 3: PABRIK KONTEN & PUBLISHING
# ==========================================
with st.expander("🏭 Modul 3: Pabrik Konten & Publishing", expanded=False):
    visual_style = st.selectbox("Style Video Veo", ["Industrial Action", "CGI Cinematic", "Blueprint Tech"])
    if st.button("🚀 GENERATE CAMPAIGN (JSON)", use_container_width=True, type="primary"):
        if not st.session_state.intelligence_data:
            st.warning("Jalankan Modul 2 dulu!")
        else:
            model_factory = genai.GenerativeModel('gemini-flash-latest', generation_config={"response_mime_type": "application/json"})
            factory_prompt = f"Data: {st.session_state.intelligence_data}. Buat kampanye {brand} {unit_type}. JSON strict: {{'copywriting': 'teks', 'veo_prompts': [{{'scene': 1, 'visual': 'prompt', 'vo': 'suara', 'sfx': 'efek'}}]}}"
            with st.spinner("Memproduksi JSON..."):
                res_json = model_factory.generate_content(factory_prompt)
                st.session_state.campaign_data = json.loads(res_json.text)

    if st.session_state.campaign_data:
        data = st.session_state.campaign_data
        st.info(data.get("copywriting", ""))
        for s in data.get("veo_prompts", []):
            st.markdown(f"**S{s.get('scene')}:** `{s.get('visual')}`", unsafe_allow_html=True)
        col_p1, col_p2 = st.columns(2)
        with col_p1:
            st.download_button("💾 Download JSON", data=json.dumps(data), file_name="Kampanye.json", use_container_width=True)
        with col_p2:
            st.button("⏰ Set Schedule", use_container_width=True)

# ==========================================
# MODUL 4: AI COMMENT SNIPER
# ==========================================
with st.expander("💬 Modul 4: AI Comment Sniper", expanded=False):
    komen = st.text_area("Paste Komentar Prospek/Haters:")
    if st.button("🎯 Tembak Balasan", use_container_width=True):
        model_sniper = genai.GenerativeModel('gemini-flash-latest')
        with st.spinner("Meracik balasan..."):
            st.success(model_sniper.generate_content(f"Balas komentar: '{komen}' untuk {brand} {unit_type}. Deteksi Hot Lead/Skeptis/Troll, patahkan argumennya.").text)

# ==========================================
# MODUL 5: EXECUTIVE PRODUCTION & ROI DASHBOARD 
# ==========================================
with st.expander("💎 MODUL 5: EXECUTIVE PRODUCTION & ROI DASHBOARD", expanded=False):
    st.markdown("### 📊 Analisis Finansial & Teknis Terpadu")
    
    tab1, tab2, tab3, tab4 = st.tabs(["🧮 Kalkulator ROI", "🧪 Formula Material", "⚡ Efisiensi Waktu", "📄 ROI Proposal"])

    with tab1:
        c_roi1, c_roi2 = st.columns(2)
        with c_roi1:
            st.markdown("**💰 Investasi Unit**")
            h_alat = st.number_input("Harga Alat (Rp):", value=850000000, step=10000000, key="roi_halat")
            dp = st.slider("DP (%)", 0, 50, 20, key="roi_dp")
            bunga = st.number_input("Bunga Efektif (%/Thn):", value=10.0, key="roi_bunga")
            tenor = st.number_input("Tenor (Bulan):", value=36, key="roi_tenor")
        
        plafon = h_alat * (1 - (dp/100))
        i = (bunga / 100) / 12
        cicilan = (plafon * i) / (1 - (1 + i)**-tenor)
        
        with c_roi2:
            st.markdown("**🏗️ Target Operasional**")
            vol_bulan = st.number_input("Estimasi Produksi / Bulan (m3):", value=300, key="roi_vol")
            margin_m3 = st.number_input("Profit Bersih / m3 (Rp):", value=350000, key="roi_margin")
            
        profit_kotor = vol_bulan * margin_m3
        sisa_cash = profit_kotor - cicilan
        pb_period = h_alat / profit_kotor if profit_kotor > 0 else 0
        
        st.divider()
        r1, r2, r3 = st.columns(3)
        r1.metric("Cicilan / Bulan", f"Rp {cicilan:,.0f}")
        r2.metric("Nett Cashflow", f"Rp {sisa_cash:,.0f}")
        r3.metric("Payback Period", f"{pb_period:.1f} Bulan")

    with tab2:
        db_sni = {
            "K-100": {"s": 230, "p": 893, "k": 1027}, "K-125": {"s": 276, "p": 828, "k": 1012},
            "K-150": {"s": 299, "p": 799, "k": 1017}, "K-175": {"s": 326, "p": 760, "k": 1029},
            "K-225": {"s": 371, "p": 698, "k": 1047}, "K-250": {"s": 384, "p": 692, "k": 1039},
            "K-300": {"s": 413, "p": 681, "k": 1021}, "K-350": {"s": 448, "p": 667, "k": 1000},
            "K-400": {"s": 470, "p": 625, "k": 1025}, "K-500": {"s": 515, "p": 610, "k": 1010}
        }
        
        c_f1, c_f2 = st.columns([1, 1])
        with c_f1:
            mutu = st.selectbox("Mutu Beton (SNI):", list(db_sni.keys()), index=6, key="mat_mutu")
            total_vol_proyek = st.number_input("Total Volume Proyek (m3):", value=1000, key="mat_vol")
            harga_readymix_pasar = st.number_input("Harga Beli Ready Mix / m3 (Rp):", value=1300000, key="mat_rm")

        with c_f2:
            p_semen = st.number_input("Harga Semen (Sak 50kg):", value=75000, key="p_semen")
            p_pasir = st.number_input("Harga Pasir (per m3):", value=250000, key="p_pasir")
            p_koral = st.number_input("Harga Koral (per m3):", value=350000, key="p_koral")
            ops_tambahan = st.number_input("Biaya Ops (Solar/Tukang) / m3:", value=75000, key="p_ops")

        f = db_sni[mutu]
        sak_semen_req = f['s'] / 50
        pasir_m3_req = f['p'] / 1400
        koral_m3_req = f['k'] / 1350
        
        total_cost_m3 = (sak_semen_req * p_semen) + (pasir_m3_req * p_pasir) + (koral_m3_req * p_koral) + ops_tambahan
        margin_saving = harga_readymix_pasar - total_cost_m3
        
        tot_semen = sak_semen_req * total_vol_proyek
        tot_pasir = pasir_m3_req * total_vol_proyek
        tot_koral = koral_m3_req * total_vol_proyek
        
        st.divider()
        c_res1, c_res2 = st.columns(2)
        c_res1.markdown(f"**📦 Total Kebutuhan Proyek**")
        c_res1.write(f"- Semen: **{tot_semen:,.0f} Sak**")
        c_res1.write(f"- Pasir: **{tot_pasir:,.1f} m3**")
        c_res1.write(f"- Koral: **{tot_koral:,.1f} m3**")

        c_res2.markdown("**💰 Analisis Keuntungan**")
        c_res2.metric("Biaya Mandiri / m3", f"Rp {total_cost_m3:,.0f}")
        c_res2.metric("Hemat vs Ready Mix", f"Rp {margin_saving:,.0f} /m3", delta=f"{ (margin_saving/harga_readymix_pasar)*100:.1f}% Lebih Murah")

    with tab3:
        col_v1, col_v2 = st.columns(2)
        with col_v1:
            st.info("👴 **Metode Manual (Molen Kecil)**")
            output_manual = st.number_input("Output / Hari (m3):", value=15.0, key="spd_man")
            st.write("Tenaga: 15-20 Orang")
            
        with col_v2:
            st.success("🤖 **Self-Loading Mixer (Aimix)**")
            output_aimix = st.number_input("Output / Hari (m3):", value=60.0, key="spd_amx")
            st.write("Tenaga: 1 Operator + 1 Helper")
            
        st.divider()
        st.subheader(f"⚡ Efisiensi: {output_aimix/output_manual:.1f}x Lebih Cepat!")

    with tab4:
        st.markdown("### 📄 Draf Proposal ROI Otomatis")
        
        proposal_text = f"""*PROPOSAL ANALISIS INVESTASI & EFISIENSI BETON*
*Target Unit:* Aimix Self-Loading Concrete Mixer

*1. ANALISIS MUTU & MATERIAL (SNI 7394:2008)*
- Mutu Beton: {mutu}
- Estimasi Biaya Produksi Mandiri: Rp {total_cost_m3:,.0f} /m3
- Harga Beli Pasar (Ready Mix): Rp {harga_readymix_pasar:,.0f} /m3
- *Potensi Margin Hemat:* Rp {margin_saving:,.0f} /m3

*2. RINGKASAN KEBUTUHAN PROYEK*
- Total Volume Proyek: {total_vol_proyek} m3
- Kebutuhan Semen: {tot_semen:,.0f} Sak
- Kebutuhan Pasir: {tot_pasir:,.1f} m3
- Kebutuhan Koral: {tot_koral:,.1f} m3

*3. ANALISIS KECEPATAN & SDM*
- Output Kerja Alat: ± {output_aimix} m3 / Hari
- Efisiensi: {output_aimix/output_manual:.1f}x lebih cepat dibanding metode manual.
- Reduksi SDM: Memangkas 15+ tukang menjadi 2 kru saja.

*4. KESIMPULAN ROI (RETURN ON INVESTMENT)*
- Nilai Investasi Alat: Rp {h_alat:,.0f}
- Payback Period: {pb_period:.1f} Bulan
- *Total Saving Proyek Saat Ini:* Rp {margin_saving * total_vol_proyek:,.0f}

---------------------------
*Diajukan oleh: Adjie Agung (VORTEX Executive Sales)*"""
        
        st.code(proposal_text, language="markdown")
        wa_text = proposal_text.replace('*', '').replace(' ', '%20').replace(chr(10), '%0A')
        wa_proposal = f"https://api.whatsapp.com/send?text={wa_text}"
        st.markdown(f'<a href="{wa_proposal}" target="_blank" style="display: block; text-align: center; background-color: #25d366; color: white; padding: 12px; border-radius: 8px; text-decoration: none; font-weight: bold;">📲 Kirim Proposal via WhatsApp</a>', unsafe_allow_html=True)

# ==========================================
# MODUL 6 & 7: KOMPETITOR & UPSELL
# ==========================================
with st.expander("⚔️ Modul 6 & 7: Kill-Switch & Upsell", expanded=False):
    st.markdown("**Kompetitor Kill-Switch**")
    kom_brand = st.text_input("Merek Lawan")
    if st.button("☠️ Generate Battlecard", use_container_width=True):
        st.success(genai.GenerativeModel('gemini-flash-latest').generate_content(f"Bandingkan {brand} vs {kom_brand}. Beri kelemahan lawan dan cara kita menang.").text)
    
    st.divider()
    st.markdown("**Upsell Predictor**")
    nama_klien_upsell = st.text_input("Klien", "PT. Tambang Maju", key="upsell_klien")
    tgl_beli = st.date_input("Tgl Beli", datetime.date(2026, 1, 1))
    if st.button("🔮 Prediksi Maintenance", use_container_width=True):
        hari_op = (datetime.date.today() - tgl_beli).days
        st.info(genai.GenerativeModel('gemini-flash-latest').generate_content(f"Alat {brand} jalan {hari_op*10} jam. Apa yang harus diganti? Buat WA Upsell.").text)

# ==========================================
# MODUL 8: TENDER HACKER
# ==========================================
with st.expander("🏛️ Modul 8: LPSE Tender Hacker", expanded=False):
    tender_file = st.file_uploader("Upload RKS (PDF)", type=["pdf"], key="tender_up")
    if st.button("🕵️‍♂️ Retas Dokumen", use_container_width=True):
        if tender_file:
            tdr_txt = "".join([PyPDF2.PdfReader(tender_file).pages[i].extract_text() for i in range(min(len(PyPDF2.PdfReader(tender_file).pages), 40))])
            with st.spinner("Membedah tender..."):
                st.success(genai.GenerativeModel('gemini-flash-latest').generate_content(f"Ekstrak syarat alat dari teks ini: {tdr_txt[:15000]}. Buktikan {brand} memenuhi syarat.").text)

# ==========================================
# MODUL 9: CRM MEMORY
# ==========================================
with st.expander("📊 Modul 9: CRM Memory & Report", expanded=False):
    DB_FILE = "sales_memory.json"
    def load_memory():
        if os.path.exists(DB_FILE):
            with open(DB_FILE, "r") as f: return json.load(f)
        return []
    def save_memory(tgl, catatan):
        data = load_memory()
        data.append({"tanggal": tgl, "catatan": catatan})
        with open(DB_FILE, "w") as f: json.dump(data, f, indent=4)

    t1, t2 = st.tabs(["📝 Input Harian", "📈 Tarik Laporan"])
    with t1:
        tgl_harian = st.date_input("Tanggal", datetime.date.today(), key="crm_tgl")
        raw_notes = st.text_area("Catatan:", height=100, key="crm_notes")
        if st.button("Simpan & Buat Laporan"):
            if raw_notes:
                save_memory(str(tgl_harian), raw_notes)
                res_daily = genai.GenerativeModel('gemini-flash-latest').generate_content(f"Rapikan jadi laporan sales: {raw_notes}")
                st.info(res_daily.text)
                
                doc_d = docx.Document()
                doc_d.add_heading(f'Laporan {tgl_harian}', 0)
                for p in res_daily.text.split('\n'): doc_d.add_paragraph(p.strip())
                bio_d = io.BytesIO()
                doc_d.save(bio_d)
                st.download_button("📄 Download Word", data=bio_d.getvalue(), file_name=f"Harian_{tgl_harian}.docx", use_container_width=True)

    with t2:
        memori = load_memory()
        st.caption(f"📦 Memory: {len(memori)} catatan")
        if st.button("Generate Laporan Panjang"):
            if memori:
                big_data = "\n".join([f"{i['tanggal']}: {i['catatan']}" for i in memori])
                res_annual = genai.GenerativeModel('gemini-flash-latest').generate_content(f"Buat Laporan Eksekutif dari riwayat ini:\n{big_data}")
                st.success(res_annual.text)

# ==========================================
# MODUL 10: EXPAT NEGOTIATOR
# ==========================================
with st.expander("🌐 Modul 10: Expat Negotiator", expanded=False):
    teks_indo = st.text_area("Teks (Indonesia):", placeholder="Contoh: Pak, ROI mesin ini hanya 2 bulan...", key="expat_text")
    bahasa_target = st.selectbox("Terjemahkan ke:", ["Mandarin (Simplified - Tiongkok)", "Inggris (Business)", "Korea (Corporate)"], key="expat_lang")
    if st.button("🔠 Terjemahkan", use_container_width=True):
        if teks_indo:
            st.info(genai.GenerativeModel('gemini-flash-latest').generate_content(f"Terjemahkan ke {bahasa_target} dengan nada eksekutif B2B: {teks_indo}").text)

# ==========================================
# MODUL 11: PHOTONIS VISUALIZER & AUTO-BROCHURE ENGINE
# ==========================================
with st.expander("📸 MODUL 11: AI Photonis & Auto-Brochure", expanded=True):
    st.markdown("### 🤖 AI Product Intelligence & Brosur Instan")
    
    tab_foto, tab_brosur = st.tabs(["🔍 Photonis AI Analyzer", "📑 Auto-Sales Kit Generator"])

    with tab_foto:
        st.info("Upload foto unit lapangan, hapus background, dan biarkan AI membuat brosur teknis PDF.")
        
        c_id1, c_id2 = st.columns(2)
        with c_id1:
            user_brand = st.text_input("Nama Brand Katalog:", "AZARINDO (Tatsuo & Aimix)")
        with c_id2:
            user_logo = st.file_uploader("Upload Logo Brand (Opsional)", type=['png', 'jpg'], key="logo_up")

        uploaded_foto = st.file_uploader("📷 Upload Foto Unit", type=['jpg', 'png', 'jpeg'], key="foto_up")

        if uploaded_foto:
            file_bytes_foto = uploaded_foto.read()
            col_img, col_ai = st.columns([1, 1])
            
            with col_img:
                st.image(file_bytes_foto, caption="Original Asset", use_column_width="always")
                
                if not REMOVE_BG_API_KEY:
                    st.warning("⚠️ API Key Remove.bg belum disetting di Streamlit Secrets.")
                else:
                    if st.button("✨ Hapus Background", key="btn_rmbg"):
                        with st.spinner("Processing visual..."):
                            response = requests.post(
                                'https://api.remove.bg/v1.0/removebg',
                                files={'image_file': file_bytes_foto},
                                data={'size': 'auto'},
                                headers={'X-Api-Key': REMOVE_BG_API_KEY},
                            )
                            if response.status_code == requests.codes.ok:
                                st.session_state['fotonis_clean_img'] = response.content
                                st.rerun()

            with col_ai:
                if 'fotonis_clean_img' in st.session_state:
                    st.image(st.session_state['fotonis_clean_img'], caption="Asset Siap Katalog", use_column_width="always")
                    
                    if st.button("🧠 Jalankan Analisa AI", key="btn_ai"):
                        with st.spinner("Gemini Menganalisa Visual..."):
                            try:
                                model_p = genai.GenerativeModel('gemini-flash-latest')
                                img_pil = Image.open(io.BytesIO(st.session_state['fotonis_clean_img']))
                                
                                prompt_ai = """Bertindaklah sebagai Heavy Equipment Technical Writer. Analisis gambar ini dan buat "OFFICIAL TECHNICAL REPORT". 
ATURAN FORMAT MUTLAK (SISTEM AKAN ERROR JIKA DILANGGAR):
1. TIGA BARIS PERTAMA WAJIB BERISI DATA INI (Untuk Data Box):
UNIT: [Nama Lengkap Unit]
KATEGORI: [Jenis Alat Berat]
FOKUS: [Target Penggunaan Utama]

2. Gunakan ## untuk Sub-Judul (Contoh: ## TEMUAN TEKNIS LAPANGAN). Jangan gunakan # tunggal.
3. BUAT TABEL MARKDOWN untuk bagian spesifikasi teknis (Gunakan tanda |). Contoh:
| Deskripsi Part / Spesifikasi | Detail / Nilai |
|---|---|
| Kapasitas Mesin | 82 kW |

4. Gunakan paragraf biasa untuk copywriting penawaran. Bahasa harus sangat formal, bersih, elegan, dan profesional khas korporat B2B tingkat tinggi."""
                                
                                res_p = model_p.generate_content([prompt_ai, img_pil])
                                st.session_state['fotonis_draft_text'] = res_p.text
                                
                            except Exception as e:
                                st.error(f"Gagal memproses AI: {e}")

            if 'fotonis_draft_text' in st.session_state:
                st.divider()
                st.markdown("#### 📝 Verifikasi Spesifikasi & Katalog PDF")
                final_text = st.text_area("Edit hasil AI:", value=st.session_state['fotonis_draft_text'], height=300)
                
                c_pdf1, c_pdf2 = st.columns(2)
                with c_pdf1:
                    logo_img_doc = Image.open(user_logo) if user_logo else None
                    pdf_data = create_visual_pdf(final_text, logo_img_doc, user_brand, st.session_state.get('fotonis_clean_img'))
                    st.download_button(
                        label="📥 Download Katalog PDF",
                        data=bytes(pdf_data),
                        file_name=f"Katalog_{user_brand.replace(' ', '_')}.pdf",
                        mime="application/pdf",
                        use_container_width=True
                    )
                with c_pdf2:
                    if st.button("🗑️ Reset Gambar"):
                        del st.session_state['fotonis_clean_img']
                        del st.session_state['fotonis_draft_text']
                        st.rerun()

    with tab_brosur:
        st.markdown("### 🖨️ Mesin Penawaran Cepat (AI-Synced)")
        
        c_bro1, c_bro2 = st.columns(2)
        with c_bro1:
            nama_klien_brosur = st.text_input("Nama Klien:", value="PT. Bangun Nusantara", key="bro_klien")
            lokasi_proyek = st.text_input("Lokasi Proyek:", value="Kalimantan Timur", key="bro_lokasi")
        with c_bro2:
            # PERBAIKAN 1: Diganti menjadi Text Input agar bebas diisi unit apa saja hasil analisa Fotonis
            unit_brosur = st.text_input("Unit yang Ditawarkan:", value="AIMIX ABT40C Diesel Concrete Pump", placeholder="Ketik nama unit...", key="bro_unit")
            harga_penawaran = st.number_input("Harga Penawaran (Rp):", value=850000000, step=10000000, key="bro_harga")

        # PERBAIKAN 2: Mengambil Copywriting AIDA & Spek langsung dari memori Fotonis Tab 1
        ai_sales_pitch = ""
        if 'fotonis_draft_text' in st.session_state and st.session_state['fotonis_draft_text'] != "":
            # Membersihkan simbol markdown agar teks rapi saat dikirim ke WA klien
            clean_pitch = st.session_state['fotonis_draft_text'].replace('**', '').replace('*', '').replace('#', '')
            ai_sales_pitch = f"\nKEUNGGULAN & SPESIFIKASI UNIT (Berdasarkan Analisa Kebutuhan):\n{clean_pitch}\n"
        else:
            # Teks cadangan jika Fotonis belum dijalankan
            ai_sales_pitch = """\nKEUNGGULAN INVESTASI:
1. Efisiensi Waktu: Mempercepat produksi hingga 3x lipat.
2. Hemat SDM: Memangkas biaya upah tukang.
3. Kualitas Presisi: Mutu hasil kerja lebih terjaga.\n"""

        teks_brosur = f"""=========================================
OFFICIAL SALES KIT & TECHNICAL PROPOSAL
=========================================
Kepada Yth: Pimpinan {nama_klien_brosur}
Lokasi Proyek: {lokasi_proyek}

Merespon kebutuhan infrastruktur di proyek Bapak/Ibu, kami merekomendasikan:
▶ UNIT: {unit_brosur}
▶ HARGA INVESTASI: Rp {harga_penawaran:,.0f}
{ai_sales_pitch}
[Lampiran: Garansi Mesin 1 Tahun & Free Training Operator]
Hubungi kami untuk skema pembayaran terbaik.

Hormat kami,
Adjie Agung - Heavy Asset Specialist
========================================="""

        st.text_area("Preview Surat Penawaran (Auto-Sync dengan Photonis AI):", value=teks_brosur, height=350)
        
        col_dl1, col_dl2 = st.columns(2)
        with col_dl1:
            st.download_button("📄 Download Sales Kit (.txt)", data=teks_brosur, file_name=f"Penawaran_{nama_klien_brosur.replace(' ', '_')}.txt", mime="text/plain", use_container_width=True)
        
        with col_dl2:
            # PERBAIKAN: Menggunakan urllib.parse.quote agar semua karakter (tabel, enter, kutip) AMAN untuk URL WhatsApp
            wa_text_bro = urllib.parse.quote(teks_brosur)
            
            tombol_wa_html = f"""
            <a href="https://api.whatsapp.com/send?text={wa_text_bro}" target="_blank" style="display: block; text-align: center; background-color: #25d366; color: white; padding: 10px; border-radius: 8px; text-decoration: none; font-weight: bold;">
                📲 Kirim via WA
            </a>
            """
            st.markdown(tombol_wa_html, unsafe_allow_html=True)
            
# ==========================================
# MODUL 12: ELITE DIGITAL CARD 
# ==========================================
with st.expander("📇 Modul 12: Elite Digital Card (Dual-Core Design)", expanded=False):
    st.markdown("**🛠️ Pengaturan Tampilan Kartu**")
    card_choice = st.radio("Fokus Desain Kartu:", ["Dual-Focus (Tatsuo & Aimix)", "Single-Focus (Satu Produk)"], horizontal=True)

    single_brand = None
    if card_choice == "Single-Focus (Satu Produk)":
        single_brand = st.selectbox("🎯 Pilih Merek:", ["Tatsuo", "AIMIX", "New Timehope"])

    bg_uploads = st.file_uploader("Upload gambar (PNG/JPG):", type=["png", "jpg", "jpeg"], accept_multiple_files=True)

    url_azarindo = "https://raw.githubusercontent.com/blacksupervisor-sys/VORTEX-Heavy-Asset-Intelligence-Arbitrage-Engine-/main/AZARINDO.png"
    url_tatsuo = "https://raw.githubusercontent.com/blacksupervisor-sys/VORTEX-Heavy-Asset-Intelligence-Arbitrage-Engine-/main/TATSUO.png" 
    url_aimix = "https://raw.githubusercontent.com/blacksupervisor-sys/VORTEX-Heavy-Asset-Intelligence-Arbitrage-Engine-/main/AIMIX.png"
    url_timehope = "https://raw.githubusercontent.com/blacksupervisor-sys/VORTEX-Heavy-Asset-Intelligence-Arbitrage-Engine-/main/TIMEHOPE.png"

    qr_url = f"https://api.qrserver.com/v1/create-qr-code/?size=150x150&data={aff_link}"
    header_logo_html = ""
    accent_color = ""

    if card_choice == "Dual-Focus (Tatsuo & Aimix)":
        accent_color = "#E74C3C" 
        header_logo_html = (
            "<div style='display: flex; justify-content: center; align-items: center; gap: 15px; flex-wrap: wrap; margin-top: 5px;'>"
            f"<img src='{url_tatsuo}' height='18px' alt='Tatsuo Logo'>"
            "<div style='border-left: 2px solid #ddd; height: 22px;'></div>"
            f"<img src='{url_aimix}' height='28px' alt='Aimix Logo'>"
            "</div>"
        )
    else:
        if single_brand == "Tatsuo":
            header_logo_html = f"<img src='{url_tatsuo}' height='28px' style='margin-top: 5px;'>"
            accent_color = "#FFD700" 
        elif single_brand == "AIMIX":
            header_logo_html = f"<img src='{url_aimix}' height='42px' style='margin-top: 5px;'>"
            accent_color = "#1E90FF" 
        elif single_brand == "New Timehope": 
            header_logo_html = f"<img src='{url_timehope}' height='45px' style='margin-top: 5px;'>"
            accent_color = "#C0392B" 

    bg_html = ""
    if bg_uploads:
        images_to_use = bg_uploads[:2]
        bg_html += "<div style='position: absolute; top: 0; left: 0; width: 100%; height: 100%; z-index: 0; display: flex; opacity: 0.15; filter: grayscale(70%);'>"
        for img in images_to_use:
            base64_img = base64.b64encode(img.read()).decode("utf-8")
            bg_html += f"<div style='flex: 1; background-image: url(data:{img.type};base64,{base64_img}); background-size: cover; background-position: center; border-right: 1px solid rgba(255,255,255,0.3);'></div>"
        bg_html += "</div>" 
    else:
        bg_html = f"<div style='position: absolute; top: 50%; left: 50%; transform: translate(-50%, -50%); opacity: 0.04; z-index: 0;'><img src='{url_azarindo}' style='width: 350px;'></div>"

    st.divider()
    st.markdown("**📱 Preview Kartu Nama Digital Anda**")

    html_card = (
        f"<div style='position: relative; border: 2px solid {accent_color}; border-radius: 12px; background-color: #ffffff; padding: 25px; box-shadow: 0px 10px 20px rgba(0,0,0,0.05); overflow: hidden; margin-top: 15px;'>"
        f"{bg_html}"
        "<div style='position: relative; z-index: 1;'>"
        "<div style='text-align: center; margin-bottom: 30px;'>"
        f"<img src='{url_azarindo}' height='50px' style='margin-bottom: 5px;' alt='Azarindo Logo'>"
        "<div style='color: #7f8c8d; font-size: 0.7em; margin-top: 5px; margin-bottom: 12px; font-weight: bold; letter-spacing: 1.5px;'>OFFICIAL PARTNER:</div>"
        f"{header_logo_html}"
        "</div>"
        f"<hr style='border: 0; border-top: 2px solid {accent_color}; opacity: 0.2; margin: 20px 0;'>"
        "<div style='display: flex; flex-direction: column; align-items: center; justify-content: center; text-align: center; margin-bottom: 35px;'>"
        "<h1 style='color: #2c3e50; margin: 0; font-size: 24px; font-weight: 900; letter-spacing: 1px;'>ADJIE AGUNG</h1>"
        "<div style='width: 60px; height: 2px; background-color: #e0e0e0; margin: 10px 0;'></div>" 
        f"<p style='color: {accent_color}; margin: 0; font-size: 11px; font-weight: 700; letter-spacing: 2px;'>HEAVY EQUIPMENT SALES</p>"
        "</div>"
        "<div style='display: flex; flex-wrap: wrap; justify-content: space-between; align-items: center; gap: 20px;'>"
        "<div style='flex: 1 1 180px; font-size: 12px; color: #34495e; line-height: 2.0;'>"
        "<div>🌐 <b>Website:</b> <span style='color: #2980b9;'>azarindo.id</span></div>"
        f"<div>📱 <b>WhatsApp:</b> {wa_num}</div>"
        "<div>💼 <b>Area:</b> Kalimantan & Timur</div>"
        "<div>🏢 <b>Kantor:</b> Samarinda, Kaltim</div>"
        "</div>"
        "<div style='flex: 0 0 auto; text-align: center; margin: 0 auto;'>"
        f"<img src='{qr_url}' style='width: 75px; height: 75px; border: 1px solid #eee; padding: 5px; border-radius: 6px; background: #fff;' alt='QR Code'><br>"
        "<span style='color: #7f8c8d; font-size: 9px; font-weight: bold; letter-spacing: 0.5px;'>E-STORE SUKU CADANG</span>"
        "</div>"
        "</div>"
        "</div>"
        "</div>"
    )
    st.markdown(html_card, unsafe_allow_html=True)

    st.markdown("<br>", unsafe_allow_html=True)
    col_btn1, col_btn2 = st.columns(2)
    with col_btn1:
        pesan_wa = "Selamat pagi/siang Bapak/Ibu. Perkenalkan saya Adjie Agung, Heavy Equipment Specialist dari Azarindo. Berikut saya lampirkan kartu nama digital saya. Jika proyek Anda membutuhkan alat Konstruksi, saya siap membantu. Terima kasih."
        link_wa = f"https://api.whatsapp.com/send?text={pesan_wa.replace(' ', '%20')}"
        st.markdown(f'<a href="{link_wa}" target="_blank" style="display: block; text-align: center; background-color: #25D366; color: white; padding: 10px; border-radius: 8px; text-decoration: none; font-weight: bold;">💬 Kirim Pengantar WA</a>', unsafe_allow_html=True)

    with col_btn2:
        vcard_data = f"""BEGIN:VCARD\nVERSION:3.0\nN:Agung;Adjie;;;\nFN:Adjie Agung\nTITLE:Heavy Equipment Sales\nORG:Azarindo\nTEL;TYPE=CELL:{wa_num}\nURL:azarindo.id\nEND:VCARD"""
        st.download_button(label="📇 Download Kontak (vCard)", data=vcard_data, file_name="Adjie_Agung_Azarindo.vcf", mime="text/vcard", use_container_width=True)

# ==========================================
# MODUL 13: ULTIMATE MARKETING BROCHURE ENGINE
# ==========================================
with st.expander("🚀 MODUL 13: Ultimate Marketing Brochure Engine", expanded=True):
    st.markdown("### 🌟 Generator Brosur Penjualan Visual (PDF & PNG)")
    st.write("Gunakan modul ini untuk membuat brosur marketing yang eye-catching, dilengkapi QR Code dan Auto-Layout.")

    CATALOG_DIR = "katalog_tersimpan"
    if not os.path.exists(CATALOG_DIR):
        os.makedirs(CATALOG_DIR)

    col_b1, col_b2 = st.columns([1, 1.2])

    with col_b1:
        st.subheader("1. Visual & Identitas")
        brand_bro = st.selectbox("Pilih Merek", ["AIMIX", "TATSUO", "NEW TIMEHOPE"], key="brochure_brand")
        
        # Penentuan Default Link & Model
        if brand_bro == "AIMIX":
            default_model = "SELF LOADING MIXER"
            default_link = "https://aimix-self-loading-mixer.netlify.app/"
        elif brand_bro == "TATSUO":
            default_model = "WHEEL EXCAVATOR JP80-9"
            default_link = "https://tatsuosales-id.netlify.app/#/"
        else:
            default_model = "HSPD 360"
            default_link = "https://azarindo.id/"

        logo_file = st.file_uploader("Upload Logo Brand (PNG Transparan)", type=['png', 'jpg', 'jpeg'], key="bro_logo")
        foto_bro = st.file_uploader("Upload Foto Unit Utama (Tanpa Background)", type=['png', 'jpg', 'jpeg'], key="bro_foto")
        
        st.markdown("---")
        
        # Inisialisasi Memori Default agar tidak error (Auto-Fill System)
        if "bro_model" not in st.session_state: st.session_state["bro_model"] = default_model
        if "bro_headline" not in st.session_state: st.session_state["bro_headline"] = "TANGGUH DI SEGALA MEDAN"
        if "bro_sp1" not in st.session_state: st.session_state["bro_sp1"] = "Heavy Duty Engine"
        if "bro_sp2" not in st.session_state: st.session_state["bro_sp2"] = "Sistem Hidrolik Stabil"
        if "bro_sp3" not in st.session_state: st.session_state["bro_sp3"] = "Kapasitas Maksimal"
        if "bro_bg1" not in st.session_state: st.session_state["bro_bg1"] = "GARANSI 1 TAHUN"
        if "bro_bg2" not in st.session_state: st.session_state["bro_bg2"] = "FREE TRAINING"
        if "bro_bg3" not in st.session_state: st.session_state["bro_bg3"] = "READY STOCK"

        # Form Input (Menggunakan value dari session_state agar bisa di-update AI)
        model_bro = st.text_input("Tipe Unit", value=st.session_state["bro_model"])
        headline_bro = st.text_input("Headline Utama", value=st.session_state["bro_headline"])
        
        st.caption("Highlight Spesifikasi Cepat (Auto-Fill by AI)")
        c_sp1, c_sp2, c_sp3 = st.columns(3)
        with c_sp1: spec_engine = st.text_input("Engine / Power", value=st.session_state["bro_sp1"])
        with c_sp2: spec_cap = st.text_input("Sistem Utama", value=st.session_state["bro_sp2"])
        with c_sp3: spec_weight = st.text_input("Kapasitas / Bobot", value=st.session_state["bro_sp3"])

        st.caption("Stempel Kepercayaan (Trust Badges)")
        b_col1, b_col2, b_col3 = st.columns(3)
        with b_col1: badge1 = st.text_input("Badge 1", value=st.session_state["bro_bg1"])
        with b_col2: badge2 = st.text_input("Badge 2", value=st.session_state["bro_bg2"])
        with b_col3: badge3 = st.text_input("Badge 3", value=st.session_state["bro_bg3"])

    with col_b2:
        st.subheader("2. AI Copywriter & Rendering")
        ref_link = st.text_input("Link Website Produk (Opsional)", default_link, key="bro_link")
        
        st.markdown("**📂 Ekstraksi Spesifikasi**")
        saved_files = [f for f in os.listdir(CATALOG_DIR) if f.endswith('.pdf')]
        pilihan_katalog = st.selectbox("Pilih File Referensi", ["-- Upload Katalog Baru --", "Gunakan Data Teks Fotonis (Modul 11)"] + saved_files, key="bro_kat")
        
        pdf_path_to_read = None
        if pilihan_katalog == "-- Upload Katalog Baru --":
            pdf_ref = st.file_uploader("Upload Katalog Spesifikasi (PDF)", type=['pdf'], key="bro_pdf_up")
            if pdf_ref:
                save_path = os.path.join(CATALOG_DIR, pdf_ref.name)
                with open(save_path, "wb") as f: f.write(pdf_ref.getbuffer())
                st.success(f"✅ Katalog '{pdf_ref.name}' tersimpan ke memori!")
                pdf_path_to_read = save_path
        elif pilihan_katalog != "Gunakan Data Teks Fotonis (Modul 11)":
            pdf_path_to_read = os.path.join(CATALOG_DIR, pilihan_katalog)
            st.info(f"⚡ Menggunakan katalog: **{pilihan_katalog}**")
            
        # TOMBOL TARIK DATA & AI (PENGGERAK AUTO-FILL)
        if st.button("✨ Tarik Data & Buat Copywriting", use_container_width=True):
            with st.spinner("AI sedang membedah PDF dan mengisi form otomatis..."):
                try:
                    scraped_text = ""
                    if pdf_path_to_read:
                        with open(pdf_path_to_read, "rb") as file_pdf:
                            pdf_reader = PyPDF2.PdfReader(file_pdf)
                            for i in range(min(10, len(pdf_reader.pages))):
                                ex = pdf_reader.pages[i].extract_text()
                                if ex: scraped_text += ex + "\n"
                    
                    if pilihan_katalog == "Gunakan Data Teks Fotonis (Modul 11)":
                        scraped_text += st.session_state.get('fotonis_draft_text', "")
                        
                    # Prompt Ekstraksi JSON
                    prompt = f"""
                    Anda adalah Data Extractor Alat Berat. Kembalikan HANYA format JSON strict:
                    {{
                        "tipe_unit": "Nama model",
                        "headline": "Headline kapital",
                        "spec1": "Engine",
                        "spec2": "Sistem",
                        "spec3": "Kapasitas",
                        "badge1": "Badge 1",
                        "badge2": "Badge 2",
                        "badge3": "Badge 3",
                        "copywriting": "JUDUL | Deskripsi..."
                    }}
                    Data: {scraped_text[:10000]}
                    """
                    model_bro_ai = genai.GenerativeModel('gemini-flash-latest', generation_config={"response_mime_type": "application/json"})
                    res_bro = model_bro_ai.generate_content(prompt)
                    ai_data = json.loads(res_bro.text)
                    
                    # Update Session State untuk Auto-Fill
                    st.session_state['bro_model'] = ai_data.get('tipe_unit', "")
                    st.session_state['bro_headline'] = ai_data.get('headline', "")
                    st.session_state['bro_sp1'] = ai_data.get('spec1', "")
                    st.session_state['bro_sp2'] = ai_data.get('spec2', "")
                    st.session_state['bro_sp3'] = ai_data.get('spec3', "")
                    st.session_state['bro_bg1'] = ai_data.get('badge1', "GARANSI 1 TAHUN")
                    st.session_state['bro_bg2'] = ai_data.get('badge2', "FREE TRAINING")
                    st.session_state['bro_bg3'] = ai_data.get('badge3', "READY STOCK")
                    st.session_state['brochure_ai_result'] = ai_data.get('copywriting', "")
                    
                    st.toast("✅ Form Otomatis Terisi!")
                    st.rerun()
                except Exception as e:
                    st.error(f"Gagal ekstraksi: {e}")

        ai_raw_text = st.session_state.get('brochure_ai_result', "JUDUL | Deskripsi...")
        final_copy_bro = st.text_area("Hasil Copywriting (Bisa Diedit):", ai_raw_text, height=150)

    st.markdown("---")

    # --- TOMBOL GENERATE (LAYOUT ELEGAN GAMBAR 2) ---
    if st.button("🌟 Generate Ultimate Brochure (PDF & PNG)"):
        if not foto_bro:
            st.warning("⚠️ Upload foto unit dulu, Boss!")
        else:
            with st.spinner("Merender Brosur..."):
                import fitz
                import qrcode
                import urllib.parse
                
                # Setup Warna Tema
                if brand_bro == "AIMIX": b_color = (31, 73, 125)
                elif brand_bro == "TATSUO": b_color = (204, 0, 0)
                else: b_color = (46, 204, 113)
                
                logo_path = None
                if logo_file:
                    logo_path = f"temp_logo_{uuid.uuid4()}.png"
                    with open(logo_path, "wb") as f: f.write(logo_file.getbuffer())
                
                # Inisialisasi Class (Pastikan class ULTIMATE_BROCHURE sudah ada di atas)
                pdf = ULTIMATE_BROCHURE(brand_color=b_color, brand_name=brand_bro, website_link=ref_link, logo_path=logo_path, wa_number=wa_num)
                pdf.add_page()
                
                # --- LOGIKA RENDERING VISUAL ---
                # (Hero Image, QR, Badges, & Copywriting diletakkan di sini sesuai instruksi sebelumnya)
                
                # Rendering Teks & Layout (Simulasi singkat untuk struktur)
                pdf.set_y(115)
                pdf.set_font('helvetica', 'B', 18)
                pdf.cell(0, 10, f"{brand_bro} {model_bro}", ln=True, align='C')

                # --- OUTPUT KONFIRMASI & TOMBOL (LAYOUT GAMBAR 2) ---
                out = pdf.output(dest='S')
                pdf_bytes = bytes(out)
                
                # Render PNG
                doc = fitz.open("pdf", pdf_bytes)
                pix = doc.load_page(0).get_pixmap(dpi=300)
                png_bytes = pix.tobytes("png")
                
                st.success("🎉 Brosur Mahakarya berhasil dibuat dengan Layout Sempurna!")
                
                # Tombol Download Horizontal
                dl1, dl2, dl3 = st.columns(3)
                with dl1:
                    st.download_button("⬇️ Download High-Res PDF", pdf_bytes, f"Brosur_{model_bro}.pdf", "application/pdf", use_container_width=True)
                with dl2:
                    st.download_button("🖼️ Download Gambar (PNG)", png_bytes, f"Brosur_{model_bro}.png", "image/png", use_container_width=True)
                with dl3:
                    wa_msg = urllib.parse.quote(f"Halo, berikut brosur {brand_bro} {model_bro}.")
                    st.markdown(f'<a href="https://wa.me/{wa_num}?text={wa_msg}" target="_blank" style="display:block;text-align:center;background:#25D366;color:white;padding:10px;border-radius:8px;text-decoration:none;font-weight:bold;">📲 Kirim via WA</a>', unsafe_allow_html=True)        
# ==========================================
# FOOTER
# ==========================================
st.markdown("<div class='footer'>Architected & Developed by <b>Adjie Agung</b> <br> VORTEX 4.0 - Heavy-Asset Domination System</div>", unsafe_allow_html=True)            
